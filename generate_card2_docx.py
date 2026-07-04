#!/usr/bin/env python3
"""Convert 卡2内容-成人.md to Word document with preserved hierarchy."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches


def set_run_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_styled_paragraph(doc, text, style=None, font_name="宋体", font_size=10.5, bold=False, alignment=None):
    paragraph = doc.add_paragraph(style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    return paragraph


def parse_markdown(md_path):
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    blocks = []
    current_section = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            blocks.append({"type": "title", "text": stripped[2:].strip()})
        elif stripped.startswith("## "):
            current_section = {"type": "section", "text": stripped[3:].strip(), "items": []}
            blocks.append(current_section)
        elif re.match(r"^\d+\.\s+", stripped) and current_section is not None:
            current_section["items"].append(stripped)

    return blocks


def build_document(blocks, output_path):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    for block in blocks:
        if block["type"] == "title":
            add_styled_paragraph(
                doc,
                block["text"],
                font_name="黑体",
                font_size=16,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            doc.add_paragraph()
            continue

        add_styled_paragraph(
            doc,
            block["text"],
            style="Heading 2",
            font_name="黑体",
            font_size=12,
            bold=True,
        )

        for item in block["items"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(item)
            set_run_font(run, "宋体", 10.5)

    doc.save(output_path)


def main():
    md_path = Path("/workspace/卡2内容-成人.md")
    output_path = Path("/workspace/卡2内容-成人.docx")
    blocks = parse_markdown(md_path)
    build_document(blocks, output_path)
    print(f"Generated: {output_path}")


if __name__ == "__main__":
    main()
