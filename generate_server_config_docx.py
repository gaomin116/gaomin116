#!/usr/bin/env python3
"""Generate hospital business server configuration and backup plan as .docx."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_font(cell, text, bold=False, size=10.5, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=14, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=12, bold=True)
    else:
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=10.5, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tblPr.append(borders)


def main():
    output_path = "/workspace/业务服务器配置与备份方案.docx"
    hardware_path = "/workspace/表1硬件环境配置表.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ===== Title =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("业务服务器清单、架构说明与备份方案")
    set_run_font(run, name="黑体", size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（慢性病管理系统 / 微信 H5 业务）")
    set_run_font(run, name="宋体", size=12)

    # ===== 1. 服务器清单 =====
    add_heading_cn(doc, "一、业务服务器清单及所需服务器资源", level=1)
    add_body(
        doc,
        "本系统采用“内网应用服务器 + DMZ 区前置机”两台虚拟机部署。"
        "应用服务器承载主程序与数据库；前置机承载 Nginx 及对外接口程序，"
        "用于开放外网访问手机端应用。",
    )

    add_heading_cn(doc, "表1  硬件环境配置表", level=2)

    servers = [
        {
            "index": "1",
            "name": "应用服务器\n（内网 / ManBing-app）",
            "config": (
                "（1）设备配置4颗处理器，2.10GHz主频（Intel Xeon Gold 5318Y）；\n"
                "（2）内存配置16GB；\n"
                "（3）硬盘配置2块VMware虚拟SCSI硬盘；\n"
                "（4）网络接口配置1块vmxnet3以太网适配器（VMware虚拟万兆网卡）；\n"
                "（5）Microsoft Windows Server 2019 Standard操作系统（64位）；"
            ),
            "count": "1台",
            "remark": (
                "部署于医院内网；VMware虚拟机；"
                "部署Tomcat应用与MariaDB数据库；"
                "计算机全名：ManBing-app.hnzyfy.com；域：hnzyfy.com"
            ),
        },
        {
            "index": "2",
            "name": "前置机\n（DMZ区）",
            "config": (
                "（1）设备配置4颗处理器，2.10GHz主频（Intel Xeon Gold 5318Y）；\n"
                "（2）内存配置8GB；\n"
                "（3）存储配置500GB；\n"
                "（4）网络接口配置于DMZ区（VMware虚拟网卡）；\n"
                "（5）Microsoft Windows Server 2019 Standard操作系统（64位）；"
            ),
            "count": "1台",
            "remark": (
                "部署于DMZ区；VMware虚拟机；"
                "部署Nginx、微信接口程序、短信接口程序；"
                "对外提供手机端H5访问入口"
            ),
        },
    ]

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    headers = ["序号", "硬件设备名称", "配置要求", "数量", "备注"]
    widths = [Cm(1.2), Cm(3.2), Cm(8.0), Cm(1.5), Cm(4.3)]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_font(
            header_cells[i],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        header_cells[i].width = widths[i]

    for server in servers:
        row_cells = table.add_row().cells
        set_cell_font(
            row_cells[0], server["index"], align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_font(
            row_cells[1], server["name"], align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_font(row_cells[2], server["config"])
        set_cell_font(
            row_cells[3], server["count"], align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_font(row_cells[4], server["remark"])
        for i, width in enumerate(widths):
            row_cells[i].width = width

    add_heading_cn(doc, "软件组件部署一览", level=2)
    soft_table = doc.add_table(rows=1, cols=4)
    soft_table.style = "Table Grid"
    soft_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(soft_table)
    soft_headers = ["服务器", "部署组件", "主要用途", "参考地址"]
    soft_widths = [Cm(3.5), Cm(4.5), Cm(5.0), Cm(4.2)]
    for i, h in enumerate(soft_headers):
        set_cell_font(
            soft_table.rows[0].cells[i],
            h,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        soft_table.rows[0].cells[i].width = soft_widths[i]

    soft_rows = [
        (
            "应用服务器（内网）",
            "Tomcat、MariaDB",
            "主程序运行与业务数据存储",
            "192.168.8.152",
        ),
        (
            "前置机（DMZ）",
            "Nginx、微信接口、短信接口",
            "外网访问入口及第三方接口转发",
            "172.17.4.87",
        ),
    ]
    for row in soft_rows:
        cells = soft_table.add_row().cells
        for i, val in enumerate(row):
            set_cell_font(
                cells[i],
                val,
                align=WD_ALIGN_PARAGRAPH.CENTER if i != 2 else None,
            )
            cells[i].width = soft_widths[i]

    # ===== 2. 架构 =====
    add_heading_cn(doc, "二、业务服务器架构说明", level=1)
    add_body(
        doc,
        "系统采用分层部署：外网用户通过医院公网IP及二级域名访问DMZ区前置机；"
        "前置机经Nginx反向代理及接口程序，单向访问内网应用服务器上的应用服务与数据库服务。"
        "内网服务器不直接对公网开放。",
    )

    add_heading_cn(doc, "2.1 架构拓扑（逻辑）", level=2)
    add_body(
        doc,
        "手机端用户 / 微信H5"
        " → 医院二级域名解析至医院公网IP"
        " → 公网IP代理转发至DMZ前置机（Nginx，端口30348）"
        " → 前置机单向访问内网应用服务器（端口3307数据库、18989 Nginx/应用）"
        " → 前置机按需访问腾讯微信接口域名（api.weixin.qq.com、mp.weixin.qq.com）。",
    )

    add_heading_cn(doc, "2.2 网络开通要求摘要", level=2)

    net_table = doc.add_table(rows=1, cols=5)
    net_table.style = "Table Grid"
    net_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(net_table)
    net_headers = ["序号", "源", "目标", "协议/端口", "用途"]
    net_widths = [Cm(1.2), Cm(4.0), Cm(4.5), Cm(3.0), Cm(4.5)]
    for i, h in enumerate(net_headers):
        set_cell_font(
            net_table.rows[0].cells[i],
            h,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        net_table.rows[0].cells[i].width = net_widths[i]

    net_rows = [
        (
            "1",
            "前置机\n172.17.4.87",
            "api.weixin.qq.com\nmp.weixin.qq.com",
            "HTTPS / TCP",
            "访问腾讯微信原生接口（单向）",
        ),
        (
            "2",
            "前置机\n172.17.4.87",
            "应用服务器\n192.168.8.152",
            "TCP 3307、18989",
            "访问数据库及Nginx/应用端口（单向）",
        ),
        (
            "3",
            "医院公网IP",
            "前置机\n172.17.4.87",
            "TCP 30348",
            "公网代理访问前置机Nginx服务（单向）",
        ),
        (
            "4",
            "医院二级域名",
            "医院公网IP",
            "DNS解析",
            "微信H5页面域名解析",
        ),
    ]
    for row in net_rows:
        cells = net_table.add_row().cells
        for i, val in enumerate(row):
            set_cell_font(
                cells[i], val, align=WD_ALIGN_PARAGRAPH.CENTER
            )
            cells[i].width = net_widths[i]

    note = doc.add_paragraph()
    run = note.add_run(
        "说明：医院公网IP及二级域名由医院侧提供；短信接口出网地址如有额外白名单要求，部署时另行确认开通。"
    )
    set_run_font(run, size=9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ===== 3. 两机房 =====
    add_heading_cn(doc, "三、两机房部署说明", level=1)
    add_body(
        doc,
        "经项目沟通确认，本期不实施两机房（双中心）部署。"
        "系统按单机房方案建设：内网应用服务器1台 + DMZ前置机1台。"
        "如后续医院有同城灾备或两机房需求，可另行评估主备切换、数据同步及网络专线条件后补充方案。",
    )
    add_bullet(doc, "本期结论：不支持/不实施两机房双活或双中心部署。")
    add_bullet(doc, "本期部署模式：单机房，虚拟化部署（VMware）。")
    add_bullet(doc, "可用性保障：依赖本机房虚拟化平台高可用能力及下文备份恢复方案。")

    # ===== 4. 备份 =====
    add_heading_cn(doc, "四、备份方案", level=1)

    add_heading_cn(doc, "4.1 备份对象与策略", level=2)
    bak_table = doc.add_table(rows=1, cols=2)
    bak_table.style = "Table Grid"
    set_table_borders(bak_table)
    bak_items = [
        ("备份对象", "MariaDB业务数据库（全量）"),
        ("备份方式", "mysqldump 逻辑全量备份"),
        ("备份频率", "全量备份（建议每日执行，具体时间窗口与医院运维协商）"),
        ("保留周期", "保留半年（约180天）"),
        ("存放位置", "应用服务器本机磁盘"),
        ("恢复责任", "由乙方运维人员负责执行恢复"),
        ("恢复说明", "根据备份文件使用mysql客户端导入恢复；恢复前确认停服窗口"),
    ]
    for k, v in bak_items:
        cells = bak_table.add_row().cells
        set_cell_font(cells[0], k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(cells[1], v)
        cells[0].width = Cm(3.5)
        cells[1].width = Cm(13.5)
    # remove empty first row created by rows=1
    tbl = bak_table._tbl
    tbl.remove(tbl.tr_lst[0])

    add_heading_cn(doc, "4.2 备份与恢复流程要点", level=2)
    add_bullet(doc, "日常：在应用服务器上按计划执行 mysqldump 全量导出，备份文件保留于本机指定目录，超期文件按半年策略清理。")
    add_bullet(doc, "异常恢复：由乙方运维定位故障点，选用对应时点备份文件完成数据库恢复，并验证业务可用性。")
    add_bullet(doc, "说明：本方案为数据库级逻辑备份；应用配置及前置机Nginx/接口配置变更后，建议同步留存配置副本。虚拟机级快照可由医院虚拟化平台按院内规范另行管理。")

    add_heading_cn(doc, "4.3 建议关注事项（供医院评审）", level=2)
    add_bullet(doc, "备份存于本机，存在与生产机同损风险；如医院有统一备份平台或异地存储，建议将 mysqldump 结果同步拷贝一份至备份平台。")
    add_bullet(doc, "应用服务器表中“2块虚拟SCSI硬盘”未标注容量，建议部署前与医院确认数据盘容量，确保可容纳半年备份文件增长。")

    # ===== 5. 汇总 =====
    add_heading_cn(doc, "五、配置结论汇总", level=1)
    add_bullet(doc, "服务器数量：2台虚拟机（内网应用服务器1台、DMZ前置机1台）。")
    add_bullet(doc, "核心资源：应用服务器 4vCPU / 16GB；前置机 4vCPU / 8GB / 500GB；操作系统均为 Windows Server 2019 Standard。")
    add_bullet(doc, "架构：前置机对外、应用服务器对内；数据库与主程序同机部署于内网。")
    add_bullet(doc, "两机房：本期不实施。")
    add_bullet(doc, "备份：mysqldump 全量备份数据库，保留半年，存本机，由乙方运维恢复。")

    doc.save(output_path)
    print(f"Generated: {output_path}")

    # Also refresh standalone 表1
    hw = Document()
    sec = hw.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    t = hw.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("表1  硬件环境配置表")
    set_run_font(r, name="黑体", size=14, bold=True)
    hw.add_paragraph()

    ht = hw.add_table(rows=1, cols=5)
    ht.style = "Table Grid"
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell_font(
            ht.rows[0].cells[i],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        ht.rows[0].cells[i].width = widths[i]
    for server in servers:
        row_cells = ht.add_row().cells
        set_cell_font(
            row_cells[0], server["index"], align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_font(
            row_cells[1], server["name"], align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_font(row_cells[2], server["config"])
        set_cell_font(
            row_cells[3], server["count"], align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_font(row_cells[4], server["remark"])
        for i, width in enumerate(widths):
            row_cells[i].width = width

    hw.save(hardware_path)
    print(f"Generated: {hardware_path}")


if __name__ == "__main__":
    main()
